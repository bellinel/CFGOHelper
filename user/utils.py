import asyncio
from docx import Document as DocxDocument
import PyPDF2
import re
import os
from aiogram import Bot
from aiogram.types import Message
import requests
from bs4 import BeautifulSoup


async def read_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    content = remove_empty_lines(content)
    return content, path

def remove_empty_lines(text):
    return "\n".join(line for line in text.splitlines() if line.strip())

async def read_docx(path):
    doc = DocxDocument(path)
    lines = []

    # Считываем абзацы
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    # Считываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    content = "\n".join(lines)
    content = remove_empty_lines(content)
    return content, path

async def read_pdf(path):
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    content = remove_empty_lines(content)
    return content, path


def remove_empty_lines(text: str) -> str:
    return "\n".join(line for line in text.splitlines() if line.strip())

def markdown_bold_to_html(text):
    return re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)



async def handle_document(msg: Message, bot: Bot):
    file_info = await bot.get_file(msg.document.file_id)
    file_path = file_info.file_path
    file_name = msg.document.file_name
    ext = file_name.split('.')[-1].lower()

    save_path = f"downloads/{file_name}"
    os.makedirs("downloads", exist_ok=True)

    await bot.download_file(file_path, destination=save_path)

    try:
        if ext == "txt":
            content = await read_txt(save_path)
        elif ext == "docx":
            content = await read_docx(save_path)
        elif ext == "pdf":
            content = await read_pdf(save_path)
        else:
            await msg.answer("❌ Поддерживаются только .txt, .docx и .pdf файлы.")
            await delete_local_file(save_path)
            return
        return content , save_path
    except Exception as e:
        await msg.answer(f"❌ Ошибка при обработке файла: {e}")
        await delete_local_file(save_path)
        return None, None
    
def save_gpt_response(name_candidate, gpt_response, date, last_number):
    text = gpt_response

    last_number +=1
    folder = "messages"
    os.makedirs(folder, exist_ok=True)

    file_path = os.path.join(folder, f"{date}_{last_number}_{name_candidate}.txt")
    name_file = f"{date}_{last_number}_{name_candidate}"
    # Сохраняем текст в файл
    with open(file_path, "a", encoding="utf-8") as f:
        f.write(text + "\n")

    return file_path, name_file



def delete_local_file(file_name):
    try:
        if os.path.exists(file_name):
            os.remove(file_name)
            print(f"Файл {file_name} удалён локально.")
        else:
            print(f"Файл {file_name} не найден.")
    except PermissionError as e:
        print(f"Ошибка удаления файла: {e}")


def remove_square_brackets(obj):
    if isinstance(obj, dict):
        return {k: remove_square_brackets(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        # Преобразуем список в строку с элементами через запятую (без скобок)
        return ', '.join(str(item) for item in obj)
    else:
        return obj
    

async def send_analize_hh_text(url,message: Message):
    a = await message.answer('Подождите минутку, я получаю текст вакансии...')
    text_vacancy = await analize_hh_url_async(url)
    if text_vacancy:
        await a.delete()
        return text_vacancy
    else:
        await a.delete()
        await message.answer('Поддерживаются только ссылки на вакансии с hh.ru')
        return None




import aiohttp
from bs4 import BeautifulSoup

async def analize_hh_url_async(url: str):
    if not url.startswith('https://hh.ru/vacancy/'):
        return None

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
    }

    async with aiohttp.ClientSession(headers=headers) as session:
        async with session.get(url) as response:
            if response.status != 200:
                return f"Ошибка {response.status}"

            html = await response.text()
            soup = BeautifulSoup(html, 'lxml')

            try:
                title = soup.find('div', class_='vacancy-title').find('h1').text.strip()
                desc_block = soup.find('div', class_='vacancy-description').find('div', class_='vacancy-section')
                if not desc_block:
                    desc_block = soup.find('div', class_='g-user-content')  # fallback
                description = desc_block.text.strip()
                lines = [line.strip() for line in description.splitlines() if line.strip()]
                cleaned_description = '\n'.join(lines)
                return f"{title}\n\n{cleaned_description}"
            except:
                return None

    
